"""
Generates project_notes.docx — a color-coded, formatted Word document
of the Recip-Ease project notes. Run with: python make_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Palette — stored as (r, g, b) tuples ──────────────────────────────────
NAVY        = (0x1B, 0x3A, 0x6B)   # section headers
BLUE        = (0x2E, 0x75, 0xB6)   # subsection headers
ORANGE      = (0xC5, 0x5A, 0x11)   # warnings / quirks label
BODY        = (0x1A, 0x1A, 0x2E)   # body text
TABLE_HDR   = (0x1B, 0x3A, 0x6B)   # table header bg
TABLE_ALT   = (0xDE, 0xEA, 0xF1)   # table alt row bg
QUIRK_BG    = (0xFF, 0xF3, 0xCD)   # callout box bg (amber)
CODE_BG     = (0xF0, 0xF0, 0xF0)   # code block bg
HISTORY_BG  = (0xE2, 0xEF, 0xDA)   # history box bg (green)
WHITE       = (0xFF, 0xFF, 0xFF)
GREEN_DARK  = (0x37, 0x5C, 0x23)   # history label text
CODE_TEXT   = (0x1F, 0x1F, 0x1F)
CODE_RED    = (0xC7, 0x25, 0x4E)
GRAY        = (0x80, 0x80, 0x80)


def rgb(t):
    """Convert (r,g,b) tuple to RGBColor."""
    return RGBColor(*t)


def hexstr(t):
    """Convert (r,g,b) tuple to uppercase hex string for XML."""
    return f"{t[0]:02X}{t[1]:02X}{t[2]:02X}"


# python-docx has no high-level API for cell/paragraph background shading or borders,
# so the next three functions write the underlying OOXML XML directly
def set_cell_bg(cell, color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hexstr(color))
    tcPr.append(shd)


def set_para_bg(para, color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hexstr(color))
    pPr.append(shd)


def set_para_borders(para, color):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"),   "thick")
    left.set(qn("w:sz"),    "24")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), hexstr(color))
    pBdr.append(left)
    pPr.append(pBdr)


def add_section_header(doc, text):
    """Navy full-width section header."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(18)
    para.paragraph_format.space_after  = Pt(4)
    run  = para.add_run(text.upper())
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = rgb(WHITE)
    set_para_bg(para, NAVY)
    para.paragraph_format.left_indent  = Cm(0.3)
    para.paragraph_format.right_indent = Cm(0.3)


def add_sub_header(doc, text):
    """Blue subsection header."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(2)
    run  = para.add_run(text)
    run.bold            = True
    run.font.size       = Pt(11.5)
    run.font.color.rgb  = rgb(BLUE)


def add_body(doc, text, space_after=4):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(space_after)
    run  = para.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = rgb(BODY)
    return para


def add_bullet(doc, text, level=0):
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_before = Pt(1)
    para.paragraph_format.space_after  = Pt(1)
    para.paragraph_format.left_indent  = Inches(0.25 + 0.2 * level)
    run  = para.add_run(text)
    run.font.size      = Pt(10.5)
    run.font.color.rgb = rgb(BODY)
    return para


def add_numbered(doc, items):
    for item in items:
        para = doc.add_paragraph(style="List Number")
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after  = Pt(2)
        run  = para.add_run(item)
        run.font.size      = Pt(10.5)
        run.font.color.rgb = rgb(BODY)


def add_code_block(doc, lines):
    """Monospaced gray block for code / file trees."""
    for line in lines:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.left_indent  = Cm(0.5)
        set_para_bg(para, CODE_BG)
        run = para.add_run(line)
        run.font.name      = "Courier New"
        run.font.size      = Pt(9)
        run.font.color.rgb = rgb(CODE_TEXT)


def add_callout(doc, label, text, bg_color=QUIRK_BG, accent_color=ORANGE):
    """Colored callout block with left border."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(6)
    para.paragraph_format.space_after  = Pt(6)
    para.paragraph_format.left_indent  = Cm(0.4)
    set_para_bg(para, bg_color)
    set_para_borders(para, accent_color)
    if label:
        r = para.add_run(f"{label}  ")
        r.bold            = True
        r.font.size       = Pt(10.5)
        r.font.color.rgb  = rgb(accent_color)
    r2 = para.add_run(text)
    r2.font.size      = Pt(10.5)
    r2.font.color.rgb = rgb(BODY)


def add_inline_code(run_text, para):
    """Helper to add inline monospace text to an existing paragraph."""
    r = para.add_run(run_text)
    r.font.name      = "Courier New"
    r.font.size      = Pt(9.5)
    r.font.color.rgb = rgb(CODE_RED)
    return r


def add_two_col_table(doc, headers, rows):
    """Two-column color-coded table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_bg(hdr_cells[i], TABLE_HDR)
        run = hdr_cells[i].paragraphs[0].runs[0]
        run.bold            = True
        run.font.color.rgb  = rgb(WHITE)
        run.font.size       = Pt(10)

    # Data rows
    for ri, row_data in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row_data):
            cells[ci].text = val
            if ri % 2 == 1:
                set_cell_bg(cells[ci], TABLE_ALT)
            for run in cells[ci].paragraphs[0].runs:
                run.font.size      = Pt(10)
                run.font.color.rgb = rgb(BODY)

    doc.add_paragraph()  # spacer


def add_qa_item(doc, question, answer):
    """Styled Q&A block."""
    qpara = doc.add_paragraph()
    qpara.paragraph_format.space_before = Pt(8)
    qpara.paragraph_format.space_after  = Pt(2)
    qrun  = qpara.add_run(f"Q: {question}")
    qrun.bold            = True
    qrun.font.size       = Pt(10.5)
    qrun.font.color.rgb  = rgb(NAVY)

    apara = doc.add_paragraph()
    apara.paragraph_format.space_before = Pt(0)
    apara.paragraph_format.space_after  = Pt(6)
    apara.paragraph_format.left_indent  = Cm(0.5)
    arun  = apara.add_run(answer)
    arun.font.size      = Pt(10.5)
    arun.font.color.rgb = rgb(BODY)


# ═══════════════════════════════════════════════════════════════════════════
# Build the document
# ═══════════════════════════════════════════════════════════════════════════
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Cover / Title ──────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(20)
title_para.paragraph_format.space_after  = Pt(4)
tr = title_para.add_run("RECIP-EASE")
tr.bold            = True
tr.font.size       = Pt(26)
tr.font.color.rgb  = rgb(NAVY)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_after = Pt(2)
sr = sub_para.add_run("Project Reference — Team Lead Edition")
sr.font.size      = Pt(13)
sr.font.color.rgb = rgb(BLUE)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.paragraph_format.space_after = Pt(20)
dr = date_para.add_run("April 2026  ·  CSCI 3020  ·  MongoDB + Flask")
dr.font.size      = Pt(10)
dr.font.color.rgb = rgb(GRAY)

doc.add_paragraph()

# ── 1. What the project is ────────────────────────────────────────────────
add_section_header(doc, "1.  What the Project Is")
add_body(doc,
    "Recip-Ease is a recipe management web app built with Python (Flask) and MongoDB Atlas. "
    "Users can browse and search recipes, write reviews, save favorites, build weekly meal plans, "
    "and get personalized recipe recommendations. The app has two parallel layers: a browser-facing "
    "HTML interface (rendered server-side with Jinja2) and a full JSON API."
)

# ── 2. Tech stack ─────────────────────────────────────────────────────────
add_section_header(doc, "2.  Tech Stack & Why")

add_sub_header(doc, "Flask (Python web framework)")
add_body(doc,
    "We originally started with FastAPI, a newer async Python framework. The problem: FastAPI is "
    "designed around async/await and an ASGI server (uvicorn), and once we added Jinja2 templates "
    "and standard synchronous database calls, the async model was getting in the way. We migrated "
    "to Flask mid-project because it's synchronous, has better beginner documentation, and its "
    "blueprint system matched how we wanted to organize the routes. "
    "You can see this in git — commits reference 'Port Terry's CRUD routes Flask -> FastAPI' "
    "and later 'Complete Flask migration' and 'Fix run.py — switch from uvicorn (ASGI) to Flask dev server.'"
)

add_sub_header(doc, "MongoDB Atlas (database)")
add_body(doc,
    "Some early commits show SQLite being used for some collections. We moved to MongoDB Atlas for the whole "
    "project because:"
)
add_bullet(doc, "Recipe data is naturally document-shaped — a recipe has an embedded list of ingredients, "
                "a meal plan has an embedded list of days. These nest directly into MongoDB documents "
                "without needing junction tables.")
add_bullet(doc, "The rubric required aggregation pipelines, which MongoDB handles well.")
add_bullet(doc, "Atlas is free-tier cloud-hosted, so the whole team connects to the same database "
                "without running anything locally.")

add_sub_header(doc, "PyMongo")
add_body(doc,
    "The official synchronous Python driver for MongoDB. We used this instead of Motor (the async driver) "
    "because we're on Flask, not FastAPI."
)

add_sub_header(doc, "Jinja2 (HTML templates)")
add_body(doc,
    "Flask's built-in templating engine. Lets us write HTML with {{ variable }} placeholders and "
    "{% for %} loops that get filled in server-side before the page is sent to the browser. Similar "
    "to Django templates or JSP if the professor comes from a Java background."
)

add_sub_header(doc, "Bootstrap 5.3 (frontend styling)")
add_body(doc, "Pre-built CSS/component library so the UI looks reasonable without spending significant time on custom CSS.")

add_sub_header(doc, "python-dotenv (environment variables)")
add_body(doc,
    "Keeps the MongoDB connection string (the URI, which contains credentials) in a .env file that "
    "is never committed to GitHub. The .env file is listed in .gitignore."
)

# ── 3. Code organization ──────────────────────────────────────────────────
add_section_header(doc, "3.  How the Code Is Organized")

add_code_block(doc, [
    "app/",
    "├── __init__.py          ← app factory, registers all blueprints",
    "├── db.py                ← one function: get_db(), returns the MongoDB database",
    "├── blueprints/",
    "│   ├── recipes.py       ← JSON API: CRUD for recipes",
    "│   ├── users.py         ← JSON API: CRUD for users",
    "│   ├── reviews.py       ← JSON API: CRUD for reviews",
    "│   ├── meal_plans.py    ← JSON API: CRUD for meal plans",
    "│   ├── views.py         ← HTML routes: home, recipe detail, forms",
    "│   ├── recommendations.py ← GET /recommendations/<user_id>",
    "│   └── dashboard.py     ← GET /dashboard (summary stats)",
    "├── templates/           ← Jinja2 HTML files",
    "└── static/              ← CSS",
    "",
    "run.py          ← starts the Flask dev server",
    "seed_db.py      ← populates all 6 collections with sample data",
    "create_indexes.py ← creates MongoDB indexes (run once after seeding)",
    "worklog.py      ← team work log entry script",
    "start.bat       ← Windows setup script (venv → install → seed → launch)",
])

doc.add_paragraph()
add_sub_header(doc, "What a Blueprint Is")
add_body(doc,
    "Flask lets you split routes across multiple files using blueprints. Each blueprint is essentially "
    "a mini-app — it registers its own routes and gets registered onto the main app in __init__.py. "
    "This is why every blueprint file starts with something like:"
)
add_code_block(doc, ["recipes_bp = Blueprint(\"recipes\", __name__)"])
doc.add_paragraph()
add_body(doc, "And __init__.py has:")
add_code_block(doc, ["app.register_blueprint(recipes_bp)"])
doc.add_paragraph()

add_sub_header(doc, "Why Blueprint Imports Are Inside create_app()")
add_callout(doc, "Key concept:",
    "If the imports were at the top of __init__.py, you'd get a circular import error: "
    "__init__.py imports from blueprints/recipes.py, which imports from app.db, which lives inside the "
    "app package, which hasn't finished loading yet. Putting the imports inside create_app() means "
    "they don't run until the function is called, by which point everything is initialized."
)

# ── 4. Database design ────────────────────────────────────────────────────
add_section_header(doc, "4.  Database Design")

add_sub_header(doc, "Collections")
add_two_col_table(doc,
    ["Collection", "What It Stores"],
    [
        ["users",        "Accounts with name, email, dietary preferences, favorite categories"],
        ["categories",   "Recipe categories (Breakfast, Lunch, Dinner, Dessert, Snack, Appetizer, Comfort)"],
        ["recipes",      "Full recipe data including an embedded ingredients array"],
        ["reviews",      "Ratings (1–5) and comments, linked to a recipe and user by ObjectId"],
        ["saved_recipes","Bookmarks — just a user_id + recipe_id pair (join collection)"],
        ["meal_plans",   "Weekly plans with an embedded days array, linked to a user"],
    ]
)

add_sub_header(doc, "Embedded vs Referenced Documents")
add_body(doc,
    "Embedded means nested data lives inside the parent document. "
    "Referenced means one document stores the ObjectId of another, linking to a separate collection."
)
add_two_col_table(doc,
    ["Pattern", "Where we used it & why"],
    [
        ["Embedded",
         "ingredients inside recipes; days inside meal_plans. "
         "You always want ingredients when you fetch a recipe, so embedding avoids a second query."],
        ["Referenced",
         "recipes store category_id and author_user_id; reviews store user_id and recipe_id. "
         "Referenced because you query these relationships independently (e.g. all recipes in a category)."],
    ]
)

add_sub_header(doc, "ObjectId — Why Every Blueprint Has _serialize()")
add_body(doc,
    "MongoDB assigns every document a unique _id field of type ObjectId automatically. "
    "ObjectId is a 12-byte binary value — JSON can't serialize it. "
    "The _serialize() function in each blueprint converts _id (and any other ObjectId fields "
    "like category_id) to plain strings before passing the document to jsonify(). "
    "It mutates the document in place, which is fine since the object isn't reused after that."
)

add_sub_header(doc, "Indexes")
add_two_col_table(doc,
    ["Index", "Why it's needed"],
    [
        ["Text index on recipe title, description, tags",
         "Makes the search bar work. $text queries silently return nothing without it."],
        ["Unique index on users.email",
         "Prevents two accounts sharing the same email."],
        ["Compound unique index on saved_recipes (user_id, recipe_id)",
         "Prevents a user from saving the same recipe twice."],
        ["Index on reviews.recipe_id",
         "Speeds up looking up all reviews for a specific recipe."],
    ]
)

# ── 5. Advanced features ──────────────────────────────────────────────────
add_section_header(doc, "5.  The Two Advanced Features")

add_sub_header(doc, "Recommendation Engine  —  GET /recommendations/<user_id>")
add_body(doc, "Located in recommendations.py. Step by step:")
add_numbered(doc, [
    "Fetch the user document. Read their dietaryPreferences and favoriteCategories.",
    "Build a $or filter — match recipes where dietaryFlags contains any of the user's dietary preferences, "
    "OR categoryId is one of their favorite categories.",
    "Run an aggregation pipeline: $match the filter → $lookup reviews → $addFields to compute average "
    "rating → $sort by rating descending → $limit to 3 results.",
    "Return the top 3 as JSON.",
])
add_callout(doc, "Edge case:",
    "$or with an empty list is a MongoDB error. If the user has no stored preferences, "
    "conditions = [] and we fall back to match_filter = {} (match everything) so new users "
    "still get results instead of a crash."
)

add_sub_header(doc, "Dashboard  —  GET /dashboard")
add_body(doc, "Located in dashboard.py. Runs three separate aggregation pipelines:")
add_two_col_table(doc,
    ["Pipeline", "What it does"],
    [
        ["Recipes per category",
         "$lookup categories → $unwind (with preserveNullAndEmptyArrays: True to keep uncategorized recipes) "
         "→ $group by category and count → $sort by count"],
        ["Top rated",
         "$lookup reviews → filter to only recipes with at least one review → "
         "$addFields avgRating → $sort → $limit 5"],
        ["Most saved",
         "$group saved_recipes by recipe → $lookup recipe title → $sort by save count → $limit 5"],
    ]
)
add_callout(doc, "Why preserveNullAndEmptyArrays?",
    "Without it, $unwind drops documents with no matching category, so the recipe count would be wrong."
)
add_callout(doc, "Why filter reviews.0 exists?",
    "$avg of an empty array returns null. That would sort unreviewed recipes into unpredictable positions "
    "in the top-rated list."
)

# ── 6. How a request flows ────────────────────────────────────────────────
add_section_header(doc, "6.  How a Request Flows Through the Code")

add_sub_header(doc, "Example: Loading a Recipe Detail Page")
add_numbered(doc, [
    "Browser sends GET /recipe/abc123",
    "Flask matches the route in views.py → recipe_detail(recipe_id)",
    "recipe_detail() calls get_db() from db.py to get the MongoDB connection",
    "Builds an aggregation pipeline: $match by _id → $lookup category, author, reviews → "
    "$unwind → $addFields avgRating — all in one query",
    "Serializes all ObjectIds to strings",
    "Calls render_template('recipe_detail.html', recipe=recipe, users=users) — "
    "Jinja2 fills in the template and returns HTML to the browser",
])

add_sub_header(doc, "Example: Creating a Recipe via JSON API")
add_numbered(doc, [
    "POST /recipes with a JSON body",
    "Flask matches the route in recipes.py → create_recipe()",
    "Validates required fields are present",
    "Converts category_id and author_user_id strings to ObjectId",
    "Defaults tags and dietary_flags to [] if missing (so $in queries don't skip this recipe)",
    "Inserts the document, returns {\"message\": ..., \"recipe_id\": ...} as JSON",
])

# ── 7. Known quirks ───────────────────────────────────────────────────────
add_section_header(doc, "7.  Known Quirks & Inconsistencies")
add_body(doc,
    "These are real things in the codebase worth knowing about if the professor looks closely."
)

add_callout(doc, "CamelCase vs snake_case field names:",
    "The API blueprints (recipes.py, users.py, etc.) use snake_case: category_id, author_user_id, "
    "dietary_preferences. However, views.py, dashboard.py, and recommendations.py still use camelCase: "
    "categoryId, authorUserId, dietaryFlags. These files were written earlier and weren't fully updated "
    "during the Flask migration. The HTML routes work because the seed data also uses camelCase internally, "
    "but the two layers aren't storing data in exactly the same format.",
    bg_color=QUIRK_BG, accent_color=ORANGE
)

add_callout(doc, "datetime.utcnow() vs datetime.now(timezone.utc):",
    "seed_db.py uses datetime.utcnow(), which produces a naive datetime (no timezone info). "
    "The API blueprints use datetime.now(timezone.utc), which is timezone-aware. Both represent UTC, "
    "but Python treats them as different types. Not a functional problem for the demo, but a real "
    "inconsistency that would matter in production.",
    bg_color=QUIRK_BG, accent_color=ORANGE
)

add_callout(doc, "New MongoClient on every get_db() call:",
    "This is intentional for simplicity. It avoids managing a shared connection object. "
    "A real production app would reuse one MongoClient instance so the connection pool is shared. "
    "For a class project hitting a small Atlas cluster, the performance impact is negligible.",
    bg_color=QUIRK_BG, accent_color=ORANGE
)

add_callout(doc, "Terysa's comment in reviews.py:",
    "The UPDATE ONE review section has a comment saying 'reviews_id, user_id and recipe_id are "
    "auto incremented primary keys.' MongoDB _ids are not auto-incremented — they're ObjectIds "
    "(binary timestamps + random bytes). The intent of the comment is correct (those fields shouldn't "
    "be updatable), just the terminology is off.",
    bg_color=QUIRK_BG, accent_color=ORANGE
)

# ── 8. Development history ────────────────────────────────────────────────
add_section_header(doc, "8.  Development History")

history_items = [
    ("Initial scaffold",
     "Flask app, Atlas connection, basic blueprint structure. app/__init__.py wired up, db.py created."),
    ("Early HTML/CSS",
     "style.css and base HTML templates added by the team."),
    ("FastAPI detour",
     "The project was briefly moved to FastAPI with uvicorn as the server. Commits reference main.py "
     "as the FastAPI entry point. This added complexity that wasn't necessary for the project scope."),
    ("Terry's routes",
     "Terry wrote CRUD routes ('Port Terry's CRUD routes Flask -> FastAPI')."),
    ("Terysa's Flask CRUD",
     "Terysa rewrote the blueprints for Flask, initially with SQLite, then updated to MongoDB. "
     "Her inline comments (what-style) document each step of the CRUD operations."),
    ("Full Flask migration",
     "Commit eac52a9 completed the migration — wired views, dashboard, and recommendations blueprints, "
     "fixed run.py to use the Flask dev server instead of uvicorn."),
    ("Cleanup",
     "Removed duplicate files, untracked ignored files, a teammate's copy of the project that had "
     "been committed by mistake."),
    ("Testing prep",
     "Added the team testing guide (TESTING.md), then start.bat to automate setup for teammates."),
    ("Comments pass",
     "Restored Terysa's inline comments from master branch (merged in). Added 'why' comments "
     "throughout explaining design decisions, edge cases, and constraints."),
]

for title, detail in history_items:
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(5)
    para.paragraph_format.space_after  = Pt(2)
    para.paragraph_format.left_indent  = Cm(0.3)
    set_para_bg(para, HISTORY_BG)
    r1 = para.add_run(f"{title}:  ")
    r1.bold            = True
    r1.font.size       = Pt(10.5)
    r1.font.color.rgb  = rgb(GREEN_DARK)
    r2 = para.add_run(detail)
    r2.font.size      = Pt(10.5)
    r2.font.color.rgb = rgb(BODY)

doc.add_paragraph()

# ── 9. Branch structure ───────────────────────────────────────────────────
add_section_header(doc, "9.  Branch Structure")

add_two_col_table(doc,
    ["Branch", "What's on it"],
    [
        ["testing-ready-4-27-2026 (current)",
         "Most complete version. Has Terysa's comments (merged from master), "
         "start.bat, updated README, and all why comments. This is the branch to demo from."],
        ["master",
         "Has Terysa's comments and an older README. Missing start.bat and the why comments. "
         "Diverged from the testing branch — both have commits the other doesn't."],
        ["main",
         "The GitHub default branch used as the PR target."],
    ]
)

# ── 10. Professor Q&A ─────────────────────────────────────────────────────
add_section_header(doc, "10.  Professor Q & A")

add_qa_item(doc,
    "Why MongoDB instead of a relational database?",
    "Recipe data is nested — a recipe has multiple ingredients, a meal plan has multiple days. "
    "In SQL you'd need junction tables and JOINs for both of those. MongoDB lets us embed those "
    "arrays directly in the document, which is simpler and avoids multiple queries. The tradeoff "
    "is that MongoDB doesn't enforce schema or relationships, so we validate types manually "
    "(converting string IDs to ObjectId, checking that 'days' is actually a list, etc.)."
)

add_qa_item(doc,
    "What's an aggregation pipeline?",
    "MongoDB's way of doing multi-step data processing in one query. Each stage transforms the "
    "data and passes the result to the next stage. Stages we used: $match (filter), $lookup "
    "(join another collection — MongoDB's equivalent of SQL JOIN), $unwind (flatten an array), "
    "$group (group and aggregate — like SQL GROUP BY), $addFields (add computed fields like avgRating), "
    "$sort, $limit, $project (shape the output — like SELECT in SQL)."
)

add_qa_item(doc,
    "How does the recommendation engine work?",
    "Fetch the user's dietary preferences and favorite categories. Build a $or filter that matches "
    "recipes satisfying either condition. Run a pipeline that joins in review data, computes average "
    "rating, sorts by rating descending, and returns the top 3. If the user has no preferences, "
    "fall back to matching all recipes so new users still get results."
)

add_qa_item(doc,
    "Why are there two kinds of routes?",
    "The JSON API routes (recipes.py, users.py, etc.) return raw JSON — meant for programmatic use "
    "or testing with tools like Postman. The HTML routes (views.py) return full rendered web pages "
    "and handle form submissions. Both hit the same MongoDB collections but serve different consumers."
)

add_qa_item(doc,
    "What does _serialize() do and why does every blueprint have one?",
    "MongoDB's ObjectId type can't be converted to JSON directly — jsonify() would throw a TypeError. "
    "_serialize() converts _id and any other ObjectId fields to plain strings before returning them. "
    "It mutates the dict in place, which is safe because we don't reuse the object after returning."
)

add_qa_item(doc,
    "Why does the text search need an index?",
    "MongoDB's $text operator only searches fields included in a text index. If you run a $text "
    "query without the index existing, you get no results and no error — it silently returns nothing. "
    "The index is created by create_indexes.py and must be run once after seeding."
)

add_qa_item(doc,
    "Why is get_db() called at the top of every route function instead of once at startup?",
    "For simplicity in a class project context — it avoids managing a shared connection object "
    "and threading concerns. PyMongo's MongoClient manages an internal connection pool, so the "
    "actual TCP connections are reused even when the client object is recreated. A production app "
    "would reuse one module-level client instance."
)

add_qa_item(doc,
    "What are the two 'advanced features' and how do they satisfy the rubric?",
    "1) Recommendation engine (GET /recommendations/<user_id>) — uses a multi-stage aggregation "
    "pipeline with $match, $lookup, $addFields, $sort, $limit to return personalized results. "
    "2) Dashboard (GET /dashboard) — three separate aggregation pipelines producing summary "
    "statistics: recipes per category, top-rated recipes, and most-saved recipes. Both demonstrate "
    "$lookup (joining collections), $group (aggregating), and $addFields (computing derived values)."
)

# ── Save ──────────────────────────────────────────────────────────────────
doc.save("project_notes.docx")
print("Saved project_notes.docx")
